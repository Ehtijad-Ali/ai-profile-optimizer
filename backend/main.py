"""
AI Profile Optimizer Pro — FastAPI Backend v2.0
• /analyze-profile  — Gemini 2.5 vision analysis
• /generate-resume  — PDF + DOCX resume generation from profile data
"""

import os, base64, json, re, logging, io
from typing import Optional
from datetime import datetime

from dotenv import load_dotenv
import google.generativeai as genai
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from PIL import Image

# ── PDF libs
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    Table, TableStyle, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# ── DOCX libs
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Setup ─────────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
log = logging.getLogger(__name__)

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
if GEMINI_API_KEY:
    log.info("GEMINI_API_KEY loaded ✓")
    genai.configure(api_key=GEMINI_API_KEY)
else:
    log.warning("GEMINI_API_KEY not set — add it to .env")

GEMINI_MODEL = "gemini-2.5-flash"

# ─── App ────────────────────────────────────────────────────────────────────────
app = FastAPI(title="AI Profile Optimizer Pro API", version="2.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["chrome-extension://<YOUR_EXTENSION_ID>", "*"],  # * for testing
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Schemas ────────────────────────────────────────────────────────────────────
class AnalyzeRequest(BaseModel):
    image: str
    mode: Optional[str] = "profile"
    page_url: Optional[str] = ""

class AnalysisResult(BaseModel):
    score: float
    strengths: list[str]
    weaknesses: list[str]
    suggestions: list[str]
    optimized_headline: str
    optimized_bio: str
    seo_keywords: list[str]
    credibility_tips: list[str]
    platform: Optional[str] = "unknown"
    mode: Optional[str] = "profile"

class ResumeData(BaseModel):
    name: Optional[str] = ""
    title: Optional[str] = ""
    email: Optional[str] = ""
    phone: Optional[str] = ""
    location: Optional[str] = ""
    profile_url: Optional[str] = ""
    headline: Optional[str] = ""
    summary: Optional[str] = ""
    skills: Optional[str] = ""
    experience: Optional[str] = ""
    education: Optional[str] = ""
    projects: Optional[str] = ""

class GenerateResumeRequest(BaseModel):
    resume_data: ResumeData
    analysis: Optional[dict] = None
    format: str = "pdf"   # "pdf" | "docx"

# ─── Helpers ────────────────────────────────────────────────────────────────────
def decode_image(b64: str, max_size=(1280, 720)):
    data = base64.b64decode(b64)
    img  = Image.open(io.BytesIO(data))
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.thumbnail(max_size, Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue(), "image/jpeg"

def parse_json(text: str) -> dict:
    text = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.MULTILINE)
    text = re.sub(r"\s*```$", "", text, flags=re.MULTILINE).strip()
    s, e = text.find("{"), text.rfind("}") + 1
    if s == -1 or e == 0:
        raise ValueError("No JSON in Gemini response")
    return json.loads(text[s:e])

def detect_platform(url: str) -> str:
    u = (url or "").lower()
    for d, p in [("linkedin.com","LinkedIn"),("github.com","GitHub"),
                 ("fiverr.com","Fiverr"),("upwork.com","Upwork"),
                 ("behance.net","Behance"),("dribbble.com","Dribbble")]:
        if d in u: return p
    return "unknown"

def build_analysis_prompt(mode: str, page_url: str) -> str:
    platform = detect_platform(page_url)
    hint = f"This profile is from **{platform}**." if platform != "unknown" else ""

    criteria_map = {
        "resume": """
- Contact information completeness
- Professional summary strength
- Work experience clarity and impact
- Skills section quality and ATS keywords
- Education presentation
- Quantified achievements
- Visual formatting and readability""",
        "rewrite": """
- Headline keyword optimization
- Bio clarity and hook
- Value proposition strength
- Niche positioning
- Call-to-action effectiveness""",
        "seo": """
- Primary keywords present
- Missing high-value keywords
- Niche specificity
- Searchability score
- Industry terminology usage""",
        "profile": """
- Profile photo professionalism
- Banner or background image quality
- Headline clarity and keyword optimization
- Bio or description strength
- Skills presentation
- Portfolio / project visibility
- Professional branding consistency
- Trustworthiness signals
- Market positioning
- Overall attractiveness to recruiters or clients"""
    }

    task_map = {
        "resume":  "Analyze this resume or CV screenshot for professional impact and ATS optimization.",
        "rewrite": "Focus specifically on the headline/title and bio/about sections. Provide optimized rewrites.",
        "seo":     "Analyze for SEO and discoverability optimization.",
        "profile": "Comprehensively analyze this professional online profile."
    }

    return f"""
You are an elite professional profile optimization expert and career coach.
{hint}

Task: {task_map.get(mode, task_map["profile"])}

Evaluate based on:
{criteria_map.get(mode, criteria_map["profile"])}

Return ONLY valid JSON (no markdown, no explanation):

{{
  "score": <0-10 with one decimal>,
  "platform": "<platform name or 'unknown'>",
  "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "weaknesses": ["<weakness 1>", "<weakness 2>", "<weakness 3>"],
  "suggestions": ["<suggestion 1>", "<suggestion 2>", "<suggestion 3>", "<suggestion 4>", "<suggestion 5>"],
  "optimized_headline": "<compelling keyword-rich headline>",
  "optimized_bio": "<3-4 sentence optimized bio>",
  "seo_keywords": ["<kw1>","<kw2>","<kw3>","<kw4>","<kw5>","<kw6>","<kw7>","<kw8>"],
  "credibility_tips": ["<tip1>","<tip2>","<tip3>","<tip4>"]
}}

Be specific. Do not invent details not visible in the screenshot.
Return ONLY the JSON object.
""".strip()

# ─── Routes: Analysis ─────────────────────────────────────────────────────────
@app.get("/")
def root():
    return {"service": "AI Profile Optimizer Pro", "version": "2.0.0", "status": "running"}

@app.get("/health")
def health():
    return {"status": "ok", "model": GEMINI_MODEL, "api_key_loaded": bool(GEMINI_API_KEY)}

@app.post("/analyze-profile", response_model=AnalysisResult)
async def analyze_profile(req: AnalyzeRequest):
    if not GEMINI_API_KEY:
        raise HTTPException(500, "GEMINI_API_KEY not set. Add it to .env file.")
    if not req.image:
        raise HTTPException(400, "No image provided.")

    try:
        img_bytes, mime = decode_image(req.image)
        log.info(f"Image: {len(img_bytes)/1024:.1f} KB | mode={req.mode}")
    except Exception as e:
        raise HTTPException(400, f"Invalid image: {e}")

    try:
        prompt = build_analysis_prompt(req.mode or "profile", req.page_url or "")
        model  = genai.GenerativeModel(GEMINI_MODEL)
        log.info("Sending to Gemini...")
        resp   = model.generate_content([{"mime_type": mime, "data": img_bytes}, prompt])
        raw    = resp.text
        log.info(f"Gemini replied ({len(raw)} chars)")
    except Exception as e:
        raise HTTPException(502, f"Gemini API error: {e}")

    try:
        result = parse_json(raw)
        score  = max(0.0, min(10.0, float(result.get("score", 5.0))))
        return AnalysisResult(
            score=score,
            platform=result.get("platform", "unknown"),
            strengths=result.get("strengths", []),
            weaknesses=result.get("weaknesses", []),
            suggestions=result.get("suggestions", []),
            optimized_headline=result.get("optimized_headline", ""),
            optimized_bio=result.get("optimized_bio", ""),
            seo_keywords=result.get("seo_keywords", []),
            credibility_tips=result.get("credibility_tips", []),
            mode=req.mode,
        )
    except Exception as e:
        log.error(f"JSON parse error: {e}\nRaw: {raw[:400]}")
        raise HTTPException(500, f"Failed to parse Gemini response: {raw[:200]}")


# ─── Routes: Resume Generation ────────────────────────────────────────────────
@app.post("/generate-resume")
async def generate_resume(req: GenerateResumeRequest):
    rd = req.resume_data
    fmt = req.format.lower()

    if fmt == "pdf":
        buf = build_pdf_resume(rd)
        media = "application/pdf"
        ext   = "pdf"
    elif fmt in ("docx", "word"):
        buf = build_docx_resume(rd)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext   = "docx"
    else:
        raise HTTPException(400, f"Unsupported format: {fmt}")

    name = (rd.name or "resume").replace(" ", "_").lower()
    filename = f"{name}_resume_{datetime.now().strftime('%Y%m%d')}.{ext}"

    return StreamingResponse(
        buf,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ─── PDF Builder ──────────────────────────────────────────────────────────────
def build_pdf_resume(rd: ResumeData) -> io.BytesIO:
    buf = io.BytesIO()

    # Accent color
    ACCENT = colors.HexColor("#4F46E5")   # indigo
    DARK   = colors.HexColor("#1E1B4B")
    GRAY   = colors.HexColor("#6B7280")
    LIGHT  = colors.HexColor("#F3F4F6")
    WHITE  = colors.white

    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.65*inch, rightMargin=0.65*inch,
        topMargin=0.55*inch,  bottomMargin=0.55*inch,
    )

    styles = getSampleStyleSheet()

    S_NAME = ParagraphStyle("S_NAME", parent=styles["Normal"],
        fontSize=24, fontName="Helvetica-Bold",
        textColor=DARK, spaceAfter=2, leading=28)

    S_TITLE = ParagraphStyle("S_TITLE", parent=styles["Normal"],
        fontSize=12, fontName="Helvetica-Oblique",
        textColor=ACCENT, spaceAfter=4)

    S_CONTACT = ParagraphStyle("S_CONTACT", parent=styles["Normal"],
        fontSize=9, fontName="Helvetica",
        textColor=GRAY, spaceAfter=8)

    S_SECTION = ParagraphStyle("S_SECTION", parent=styles["Normal"],
        fontSize=10, fontName="Helvetica-Bold",
        textColor=ACCENT, spaceBefore=10, spaceAfter=3,
        borderPadding=(0,0,2,0), textTransform="uppercase", letterSpacing=0.8)

    S_BODY = ParagraphStyle("S_BODY", parent=styles["Normal"],
        fontSize=10, fontName="Helvetica",
        textColor=DARK, leading=15, spaceAfter=4)

    S_BULLET = ParagraphStyle("S_BULLET", parent=styles["Normal"],
        fontSize=10, fontName="Helvetica",
        textColor=DARK, leading=14, leftIndent=14, spaceAfter=3)

    story = []

    # ── Header ──
    name_str  = rd.name or "Your Name"
    title_str = rd.title or rd.headline or ""
    contacts  = [c for c in [rd.email, rd.phone, rd.location, rd.profile_url] if c]

    story.append(Paragraph(name_str, S_NAME))
    if title_str:
        story.append(Paragraph(title_str, S_TITLE))
    if contacts:
        story.append(Paragraph(" · ".join(contacts), S_CONTACT))

    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=8))

    def section(title, content_paragraphs):
        if not content_paragraphs: return
        story.append(Paragraph(title, S_SECTION))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E5E7EB"), spaceAfter=5))
        for p in content_paragraphs:
            story.append(p)

    # ── Summary ──
    if rd.summary:
        section("Professional Summary", [Paragraph(rd.summary, S_BODY)])

    # ── Skills ──
    if rd.skills:
        skill_items = [s.strip() for s in rd.skills.split(",") if s.strip()]
        if skill_items:
            # Build a 3-column chip-style table
            COLS = 3
            rows = [skill_items[i:i+COLS] for i in range(0, len(skill_items), COLS)]
            # Pad last row
            while len(rows[-1]) < COLS:
                rows[-1].append("")

            tdata = [[Paragraph(cell, ParagraphStyle("chip", fontSize=9, fontName="Helvetica",
                                textColor=ACCENT)) for cell in row] for row in rows]

            tw = 7.2 * inch
            t  = Table(tdata, colWidths=[tw/COLS]*COLS)
            t.setStyle(TableStyle([
                ("BACKGROUND",  (0,0), (-1,-1), colors.HexColor("#EEF2FF")),
                ("ROUNDEDCORNERS", [4]),
                ("LEFTPADDING",  (0,0), (-1,-1), 8),
                ("RIGHTPADDING", (0,0), (-1,-1), 8),
                ("TOPPADDING",   (0,0), (-1,-1), 5),
                ("BOTTOMPADDING",(0,0), (-1,-1), 5),
                ("ROWBACKGROUNDS",(0,0),(-1,-1),[colors.HexColor("#EEF2FF")]),
                ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#C7D2FE")),
            ]))
            section("Skills", [t, Spacer(1,4)])

    # ── Experience ──
    if rd.experience:
        lines = [l for l in rd.experience.split("\n") if l.strip()]
        paras = []
        for line in lines:
            clean = line.strip().lstrip("•-").strip()
            paras.append(Paragraph(f"• {clean}", S_BULLET))
        section("Work Experience", paras)

    # ── Education ──
    if rd.education:
        lines = [l for l in rd.education.split("\n") if l.strip()]
        paras = [Paragraph(l.strip(), S_BODY) for l in lines]
        section("Education", paras)

    # ── Projects ──
    if rd.projects:
        lines = [l for l in rd.projects.split("\n") if l.strip()]
        paras = []
        for line in lines:
            clean = line.strip().lstrip("•-").strip()
            paras.append(Paragraph(f"• {clean}", S_BULLET))
        section("Projects & Portfolio", paras)

    doc.build(story)
    buf.seek(0)
    return buf


# ─── DOCX Builder ─────────────────────────────────────────────────────────────
def build_docx_resume(rd: ResumeData) -> io.BytesIO:
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    ACCENT_RGB = RGBColor(0x4F, 0x46, 0xE5)   # indigo
    DARK_RGB   = RGBColor(0x1E, 0x1B, 0x4B)
    GRAY_RGB   = RGBColor(0x6B, 0x72, 0x80)

    def set_font(run, size, bold=False, italic=False, color=None):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def add_hr(doc, color_hex="4F46E5"):
        p    = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_title(doc, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(title.upper())
        set_font(run, 10, bold=True, color=ACCENT_RGB)
        run.font.name = "Calibri"
        add_hr(doc)

    # ── Name ──
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after  = Pt(2)
    nr = name_para.add_run(rd.name or "Your Name")
    set_font(nr, 22, bold=True, color=DARK_RGB)

    # ── Headline / Title ──
    if rd.headline or rd.title:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_after = Pt(3)
        hr = hp.add_run(rd.headline or rd.title or "")
        set_font(hr, 12, italic=True, color=ACCENT_RGB)

    # ── Contact line ──
    contacts = [c for c in [rd.email, rd.phone, rd.location, rd.profile_url] if c]
    if contacts:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(6)
        cr = cp.add_run(" · ".join(contacts))
        set_font(cr, 9, color=GRAY_RGB)

    add_hr(doc, "4F46E5")

    # ── Summary ──
    if rd.summary:
        section_title(doc, "Professional Summary")
        sp = doc.add_paragraph(rd.summary)
        sp.paragraph_format.space_after = Pt(4)
        for run in sp.runs:
            set_font(run, 10, color=DARK_RGB)

    # ── Skills ──
    if rd.skills:
        skills_list = [s.strip() for s in rd.skills.split(",") if s.strip()]
        if skills_list:
            section_title(doc, "Skills")
            # Build a table
            COLS = 3
            rows = [skills_list[i:i+COLS] for i in range(0, len(skills_list), COLS)]
            while len(rows[-1]) < COLS:
                rows[-1].append("")

            tbl = doc.add_table(rows=len(rows), cols=COLS)
            tbl.style = "Table Grid"
            for r_idx, row_data in enumerate(rows):
                row = tbl.rows[r_idx]
                for c_idx, cell_text in enumerate(row_data):
                    cell = row.cells[c_idx]
                    cell.text = cell_text
                    # Style the cell
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_font(run, 9, color=ACCENT_RGB)
                        # Cell background
                        tc_pr = cell._tc.get_or_add_tcPr()
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:fill"), "EEF2FF")
                        shd.set(qn("w:val"),  "clear")
                        tc_pr.append(shd)

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Experience ──
    if rd.experience:
        section_title(doc, "Work Experience")
        for line in rd.experience.split("\n"):
            line = line.strip().lstrip("•-").strip()
            if not line: continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            set_font(run, 10, color=DARK_RGB)

    # ── Education ──
    if rd.education:
        section_title(doc, "Education")
        for line in rd.education.split("\n"):
            if not line.strip(): continue
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                set_font(run, 10, color=DARK_RGB)

    # ── Projects ──
    if rd.projects:
        section_title(doc, "Projects & Portfolio")
        for line in rd.projects.split("\n"):
            line = line.strip().lstrip("•-").strip()
            if not line: continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            set_font(run, 10, color=DARK_RGB)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
